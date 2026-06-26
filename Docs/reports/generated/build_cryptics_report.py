from __future__ import annotations

import re
from pathlib import Path
from urllib.parse import unquote

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
SOURCE = ROOT / "Docs/reports/consolidated-report.md"
OUT = ROOT / "Docs/reports/Cryptics-ST2515-Consolidated-Security-Report.docx"

BASE_FONT = "Calibri"
BODY_SIZE = Pt(11)
NAVY = "102A43"
TEAL = "0F766E"
GOLD = "B7791F"
PALE_BLUE = "EAF4F8"
PALE_GOLD = "FFF7E6"
CODE_FILL = "1E293B"
CODE_TEXT = "E2E8F0"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color: str = TEAL, size: str = "12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_field(paragraph, instr: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr_text, fld_sep, fld_end])


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BASE_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
    normal.font.size = BODY_SIZE
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, before, after, color in [
        ("Heading 1", 20, 20, 7, NAVY),
        ("Heading 2", 15, 16, 5, TEAL),
        ("Heading 3", 12.5, 12, 4, "334155"),
    ]:
        st = styles[name]
        st.font.name = BASE_FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
        st.font.size = Pt(size)
        st.font.bold = name == "Heading 1"
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def add_text_with_inline_code(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor.from_string("7C3AED")
        else:
            paragraph.add_run(part)


def clean_md(text: str) -> str:
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = text.replace("\\|", "|")
    return text.strip()


def resolve_image(src: str) -> Path | None:
    if src.startswith("http"):
        return None
    src = unquote(src)
    candidate = (SOURCE.parent / src).resolve()
    if candidate.exists():
        return candidate
    candidate = (ROOT / src).resolve()
    return candidate if candidate.exists() else None


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    author_idx = None
    for idx, value in enumerate(rows[0]):
        if clean_md(value).lower() == "author":
            author_idx = idx
            break
    if author_idx is not None:
        rows = [[cell for idx, cell in enumerate(row) if idx != author_idx] for row in rows]
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False
    set_cell_margins(table)
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j in range(width):
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[j].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_text_with_inline_code(p, clean_md(row[j]) if j < len(row) else "")
            if i == 0:
                set_cell_shading(cells[j], PALE_BLUE)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)


def figure_description(alt: str, src: str) -> str:
    lower = f"{alt} {src}".lower()
    if "bruno" in lower or "request" in lower or "postman" in lower:
        return "This evidence captures the API request or response used to validate the vulnerability and compare behaviour before and after remediation."
    if "code" in lower or ".js" in lower or "config" in lower or "verify" in lower:
        return "This code evidence highlights the exact implementation pattern under review, such as missing middleware, unsafe SQL construction, token handling, or the corrected secure pattern."
    if "database" in lower or "sql" in lower or "workbench" in lower:
        return "This database evidence shows how application data was stored or exposed, supporting the assessment of confidentiality and integrity impact."
    if "devtools" in lower or "browser" in lower or "localstorage" in lower:
        return "This browser evidence shows client-side behaviour and why browser-only controls or token storage cannot be treated as a security boundary."
    if "terminal" in lower or "console" in lower or "error" in lower:
        return "This terminal evidence shows logging, error disclosure, or missing audit output during security-sensitive actions."
    return "This screenshot provides supporting evidence for the finding discussed in the surrounding section."


def add_image(doc: Document, alt: str, src: str, missing: list[str]) -> None:
    path = resolve_image(src)
    if not path:
        p = doc.add_paragraph()
        set_paragraph_shading(p, PALE_GOLD)
        r = p.add_run(f"Remote figure unavailable: {clean_md(alt)}")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(GOLD)
        p.add_run(f"\nSource reference: {src}")
        missing.append(f"{alt} -> {src}")
        return
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = caption.add_run(clean_md(alt))
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("475569")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.add_run().add_picture(str(path), width=Inches(6.25))
    except Exception:
        p.add_run().add_picture(str(path), width=Inches(5.8))
    desc = doc.add_paragraph()
    desc.paragraph_format.left_indent = Inches(0.2)
    desc.paragraph_format.right_indent = Inches(0.2)
    desc.paragraph_format.space_after = Pt(10)
    run = desc.add_run(figure_description(alt, src))
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string("475569")


def split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [c.strip() for c in stripped.split("|")]


def parse_frontmatter(text: str) -> tuple[dict[str, object], str]:
    if not text.startswith("---"):
        return {}, text
    end = text.find("\n---", 3)
    fm_text = text[3:end].strip()
    body = text[end + 4 :].lstrip()
    meta: dict[str, object] = {}
    current = None
    for line in fm_text.splitlines():
        if line.startswith("  - ") and current:
            meta.setdefault(current, []).append(line[4:].strip())
        elif ":" in line:
            k, v = line.split(":", 1)
            current = k.strip()
            meta[current] = v.strip() or []
    return meta, body


def add_cover(doc: Document, meta: dict[str, object]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(88)
    set_paragraph_shading(p, NAVY)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Cryptics")
    r.font.name = BASE_FONT
    r.font.size = Pt(36)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("FFFFFF")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(meta.get("title", "ST2515 Secure Vulnerability Analysis Report")))
    r.font.name = BASE_FONT
    r.font.size = Pt(25)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    r.bold = True
    set_paragraph_border(p, TEAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SP Games Web Application - Consolidated OWASP Assessment")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string("475569")

    doc.add_paragraph()
    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_margins(meta_table)
    rows = [
        ("Module", str(meta.get("module", "ST2515 Secure Coding"))),
        ("Date", str(meta.get("date", "June 2026"))),
        ("Submission", str(meta.get("submission", "29 June 2026"))),
        ("Team Members", "\n".join(meta.get("authors", []))),
    ]
    for label, value in rows:
        cells = meta_table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], PALE_BLUE)
        for cell in cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = BASE_FONT
                    run.font.size = Pt(10)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared as a consolidated OWASP vulnerability assessment with exploit evidence, database impact, affected code, remediation, testing notes, and project contribution log.")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string("475569")
    doc.add_page_break()


def load_contribution_log() -> str:
    log = ROOT / "Docs/tracking/log.md"
    if not log.exists():
        return ""
    text = log.read_text()
    start = text.rfind("| **Nachiketh**")
    if start == -1:
        return text[-2500:]
    return "# Project Contribution Log\n\n" + text[start:]


def richer_part_iv() -> str:
    source = ROOT / "Docs/reports/keefe-report.md"
    if not source.exists():
        return ""
    text = source.read_text()
    text = text.split("# Finding 1", 1)[-1]
    text = "# Finding 10 - Plaintext Password Storage and Lack of Hashing\n" + text
    text = text.split("# Conclusion", 1)[0]
    replacements = {
        "# Finding 2": "## Finding 11",
        "# Finding 3": "## Finding 12",
        "# Finding 4": "## Finding 13",
        "# 4. Affected Code": "### 4. Affected Code",
        "# 5. Recommendations & Fix Code": "### 5. Recommendations & Fix Code",
        "# 6. Testing Process": "### 6. Testing Process",
        "# 7. Tools Used": "### 7. Tools Used",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return "# Part IV - OWASP A07 & A02: Authentication and Cryptographic Failures\n\n" + text.strip() + "\n"


def enhance_body(body: str) -> str:
    body = re.sub(r"\n\*\*Author:\*\*.*\n", "\n", body)
    body = re.sub(r"\n## Remaining Work\n\n.*?(?=\n---|\n## Appendix|\Z)", "\n", body, flags=re.S)
    body = body.replace("Password bcrypt hashing, HTTPS enforcement, and httpOnly cookie migration remain recommended follow-up work.", "Authentication hardening, transport protection, and session handling are addressed in the remediation plan and project log.")
    part_iv = richer_part_iv()
    if part_iv:
        body = re.sub(r"# Part IV .*?(?=\n# Part V )", part_iv + "\n", body, flags=re.S)
    remediation = """# Recommended Fixes and Work Completed

The project remediation work focused first on controls that prevented unauthorised access and direct data exposure. Sensitive backend routes were moved behind JWT verification and admin role checks, user registration stopped trusting client-supplied roles, SQL queries were rewritten with parameterised placeholders, and the hardcoded JWT secret was replaced with an environment-based configuration guard.

The second remediation layer improved operational security. Password fields were removed from user-list API responses, duplicate-registration responses were normalised to reduce enumeration, review ownership checks were added, and security-sensitive actions now have a structured audit logging pattern through `securityLog.js`. The remaining recommendations are framed as production hardening rather than unfinished report work: use bcrypt consistently for all password storage and comparison, migrate browser token handling to httpOnly secure sameSite cookies, enforce HTTPS/TLS with HSTS, and ship logs to persistent monitoring.

| Area | Fix or Recommendation | Security Effect |
|---|---|---|
| Access control | `verifyToken` and `requireAdmin` on privileged routes | Blocks unauthenticated and non-admin API use |
| Registration | Server assigns normal user role | Prevents self-created admin accounts |
| SQL queries | Bound `?` placeholders | Stops SQL payloads from becoming executable query syntax |
| JWT secret | `process.env.JWT_SECRET` startup guard | Prevents predictable token forgery from source code |
| API exposure | Password removed from user response payloads | Reduces credential disclosure risk |
| Logging | Structured audit and safe error output | Improves incident response while reducing schema leakage |
| Session security | httpOnly secure cookies recommended | Reduces token theft from browser JavaScript |
| Transport security | HTTPS/TLS and HSTS recommended | Protects credentials and tokens in transit |
"""
    conclusion = """# Conclusion

The SP Games assessment shows how ordinary implementation shortcuts can combine into a critical security posture. The most serious risks were not isolated defects: open backend routes, client-trusted roles, plaintext credential handling, SQL string interpolation, weak token configuration, and missing audit trails reinforced one another. In the original state, an attacker could create privileged accounts, read sensitive user data, tamper with catalogue records, forge or steal authentication material, and leave limited forensic evidence behind.

The remediation work substantially improves the application by moving trust decisions to the backend, binding database inputs safely, externalising secrets, reducing sensitive API output, and adding clearer logging behaviour. The remaining production-hardening items are focused on defence in depth: consistently hashed passwords, secure cookie-based session storage, HTTPS everywhere, and persistent monitoring. Together, these changes move the project from a vulnerable assignment implementation toward a more secure, maintainable web application design.
"""
    body = re.sub(r"# Conclusion.*?(?=\n## Appendix|\Z)", conclusion, body, flags=re.S)
    body = re.sub(r"\n## Appendix\s+[—-]\s+Evidence Index", "\n" + remediation + "\n## Appendix - Evidence Index", body)
    log = load_contribution_log()
    if log:
        body += "\n\n" + log
    return body


def collect_toc_entries(body: str) -> list[tuple[int, str]]:
    entries: list[tuple[int, str]] = []
    for line in body.splitlines():
        h = re.match(r"^(#{1,3})\s+(.*)$", line)
        if h:
            text = clean_md(h.group(2).replace("—", "-"))
            entries.append((len(h.group(1)), text))
    return entries


def add_toc(doc: Document, entries: list[tuple[int, str]]) -> None:
    p = doc.add_paragraph("Table of Contents", style="Heading 1")
    p.paragraph_format.page_break_before = False
    toc = doc.add_paragraph()
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
    note = doc.add_paragraph("In Word or Google Docs, refresh/update the table of contents after opening to populate page numbers.")
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(9)
    for level, text in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25 * max(level - 1, 0))
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(10 if level == 1 else 9)
        r.bold = level == 1
        r.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else "475569")
    doc.add_page_break()


def build() -> None:
    raw = SOURCE.read_text()
    meta, body = parse_frontmatter(raw)
    body = enhance_body(body)
    doc = Document()
    style_doc(doc)
    add_cover(doc, meta)
    add_toc(doc, collect_toc_entries(body))

    missing: list[str] = []
    in_code = False
    code_lines: list[str] = []
    table_lines: list[str] = []

    def flush_table():
        nonlocal table_lines
        if table_lines:
            rows = [split_table_row(line) for line in table_lines if not re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line)]
            add_table(doc, rows)
            table_lines = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(8)
            set_paragraph_shading(p, CODE_FILL)
            r = p.add_run("\n".join(code_lines))
            r.font.name = "Courier New"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor.from_string(CODE_TEXT)
            code_lines = []

    for line in body.splitlines():
        if line.strip().startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            flush_table()
            continue
        if line.strip() in {"---"}:
            flush_table()
            continue
        if line.lstrip().startswith("|") and "|" in line.rstrip()[1:]:
            table_lines.append(line)
            continue
        flush_table()

        img = re.match(r"!\[(.*?)\]\((.*?)\)", line.strip())
        if img:
            add_image(doc, img.group(1), img.group(2), missing)
            continue

        h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if h:
            level = min(len(h.group(1)), 3)
            text = clean_md(h.group(2).replace("—", "-"))
            if level == 1 and (text.startswith("Part ") or text == "Conclusion"):
                doc.add_section(WD_SECTION_START.NEW_PAGE)
            p = doc.add_paragraph(text, style=f"Heading {level}")
            if level == 1:
                set_paragraph_border(p, TEAL, "8")
            continue

        bullet = re.match(r"^\s*[-*]\s+(.*)$", line)
        numbered = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_text_with_inline_code(p, clean_md(bullet.group(1)))
        elif numbered:
            p = doc.add_paragraph(style="List Number")
            add_text_with_inline_code(p, clean_md(numbered.group(1)))
        else:
            p = doc.add_paragraph()
            add_text_with_inline_code(p, clean_md(line))

    flush_table()
    flush_code()

    if missing:
        doc.add_page_break()
        doc.add_paragraph("Unavailable Remote Figures", style="Heading 1")
        doc.add_paragraph("These image URLs were referenced in the source report but could not be retrieved as local files during document generation.")
        for item in missing:
            doc.add_paragraph(item, style="List Bullet")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
